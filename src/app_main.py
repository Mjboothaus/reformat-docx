import streamlit as st
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Mm


def make_a4(document):
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)


def set_margins(document, margin_mm):
    for section in document.sections:
        section.left_margin = Mm(margin_mm)
        section.right_margin = Mm(margin_mm)

def set_top_bottom_margins(document, margin_mm):
    for section in document.sections:
        # section.top_margin = Mm(margin_mm)
        section.bottom_margin = Mm(margin_mm)

def get_margins(document):
    for section in document.sections:
        st.write((section.top_margin/914400) * 25.4, (section.bottom_margin / 914400 * 25.4))


st.header("Word doc re-formatter")
margin_mm = st.sidebar.slider(
    "Page margin width (mm):", min_value=4, max_value=20, value=8)

docx_file = st.file_uploader("Choose Word document (.docx) to convert", 'docx')

convert_button = st.button("Convert document")

if convert_button and docx_file is not None:
    docx_filename = docx_file.name
    document = Document(docx_file)

    # get_margins(document)

    make_a4(document)
    set_margins(document, margin_mm)
    set_top_bottom_margins(document, 10.583)

    for table in document.tables:
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            #     row_data = []
            for cell in row.cells:
                cell.autofit = True
        #         row_data.extend(paragraph.text for paragraph in cell.paragraphs)
        #     st.write("\t".join(row_data))
    new_filename = "output/" + docx_filename.replace(".docx", "_modified.docx")
    if Path(new_filename).exists():
        Path(new_filename).unlink()
    st.info(f"Saving: {new_filename}")
    document.save(new_filename)
    st.info("Finished")
